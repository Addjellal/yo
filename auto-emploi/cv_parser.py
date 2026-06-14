import base64
import hashlib
import os
import re
import sys
import tempfile
from pathlib import Path


def _cache_path(path: Path) -> Path:
    from config import config
    cache_dir = Path(config.output_dir).resolve()
    cache_dir.mkdir(parents=True, exist_ok=True)
    # Hash sur le chemin absolu pour éviter les collisions entre CVs
    # de même nom dans des dossiers différents. Le stem n'est que cosmétique :
    # on le neutralise strictement ([A-Za-z0-9_-]) pour qu'aucun « ../ » ni
    # séparateur ne puisse faire sortir l'écriture du cache de output/.
    digest = hashlib.sha1(str(path.resolve()).encode("utf-8")).hexdigest()[:10]
    safe_stem = re.sub(r"[^A-Za-z0-9_-]", "_", path.stem)[:50] or "cv"
    return cache_dir / f".cv_{safe_stem}_{digest}.txt"


def _read_text_tolerant(path: Path) -> str:
    """Lit un fichier texte sans crasher sur un encodage non-UTF-8 (CV exportés
    en latin-1/cp1252, ou avec BOM) : UTF-8 d'abord, repli cp1252, puis
    remplacement des octets restants."""
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", "replace")


def _write_cache_atomic(cache: Path, text: str) -> None:
    """Écrit le cache du CV (donnée personnelle) en 0600 dès la création, via
    fichier temporaire + remplacement atomique — pas de fenêtre à 0644."""
    fd, tmp = tempfile.mkstemp(dir=str(cache.parent), prefix=".cvtmp_")
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as fh:
            fh.write(text)
        if sys.platform != "win32":
            try:
                os.chmod(tmp, 0o600)
            except OSError:
                pass
        os.replace(tmp, cache)
    except Exception:
        try:
            os.unlink(tmp)
        except OSError:
            pass
        raise


_MAX_CV_BYTES = 25 * 1024 * 1024  # garde-fou : un CV ne devrait jamais dépasser 25 Mo


def parse_cv(cv_path: str, force_reparse: bool = False) -> str:
    path = Path(cv_path)
    if not path.is_file():
        raise FileNotFoundError(f"CV introuvable : {cv_path}")
    if path.stat().st_size > _MAX_CV_BYTES:
        raise ValueError(f"Fichier trop volumineux ({path.stat().st_size // 1024 // 1024} Mo, max 25 Mo).")

    # Utiliser le cache si disponible et plus récent que le PDF
    cache = _cache_path(path)
    if not force_reparse and cache.exists() and cache.stat().st_mtime >= path.stat().st_mtime:
        from rich.markup import escape
        from app_utils import console
        console.print(f"[dim]CV chargé depuis le cache ({escape(cache.name)})[/dim]")
        return cache.read_text(encoding="utf-8")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _parse_pdf(path)
    elif suffix == ".docx":
        text = _parse_docx(path)
    elif suffix == ".txt":
        text = _read_text_tolerant(path)
    else:
        raise ValueError(f"Format non supporté : {suffix}. Utilisez PDF, DOCX ou TXT.")

    # Sauvegarder en cache — le CV est une donnée personnelle : accès propriétaire
    # seul, écriture atomique en 0600 (aucune fenêtre à 0644).
    _write_cache_atomic(cache, text)
    return text


def _parse_pdf(path: Path) -> str:
    """Try native text extraction first; fall back to Claude Vision OCR for image-based PDFs."""
    text = _extract_text_native(path)
    if text:
        return text
    # PDF is image-based (scanned/rasterised) — use Claude Vision
    return _extract_text_vision(path)


def _extract_text_native(path: Path) -> str:
    """Extract selectable text using PyMuPDF (fast, no API call)."""
    try:
        import fitz  # pymupdf
        doc = fitz.open(str(path))
        pages = [page.get_text() for page in doc]
        return "\n".join(p for p in pages if p.strip())
    except ImportError:
        pass

    # Fallback: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(p for p in pages if p.strip())
    except ImportError:
        pass

    return ""


def _extract_text_vision(path: Path) -> str:
    """Render each PDF page as image and OCR via Claude Vision or Ollama vision."""
    try:
        import fitz
    except ImportError:
        raise ValueError(
            "PDF image-based détecté mais PyMuPDF non installé.\n"
            "Installez-le : pip install pymupdf"
        )

    from config import config
    from app_utils import console

    doc = fitz.open(str(path))
    console.print(f"[yellow]PDF image-based détecté ({len(doc)} page(s)) — OCR via {config.provider}...[/yellow]")

    all_text: list[str] = []
    ocr_fn = _ocr_page_ollama if config.provider == "ollama" else _ocr_page_anthropic

    for page_num, page in enumerate(doc):
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("png")
        page_text = ocr_fn(img_bytes, config)
        if page_text:
            all_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
        console.print(f"[dim]  Page {page_num + 1}/{len(doc)} traitée[/dim]")

    if not all_text:
        raise ValueError("Impossible d'extraire le texte du CV, même via OCR.")

    return "\n\n".join(all_text)


_OCR_PROMPT = (
    "Voici la page d'un CV. Extrait tout le texte visible "
    "en préservant la structure (sections, puces, dates). "
    "Réponds uniquement avec le texte extrait, sans commentaire."
)


def _ocr_page_anthropic(img_bytes: bytes, config) -> str:
    import anthropic
    if not config.anthropic_api_key:
        raise ValueError("ANTHROPIC_API_KEY manquante — impossible de faire l'OCR.")
    client = anthropic.Anthropic(api_key=config.anthropic_api_key, max_retries=3)
    img_b64 = base64.standard_b64encode(img_bytes).decode()
    response = client.messages.create(
        model=config.anthropic_model,
        max_tokens=4096,
        messages=[{
            "role": "user",
            "content": [
                {"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": img_b64}},
                {"type": "text", "text": _OCR_PROMPT},
            ],
        }],
    )
    return next((b.text for b in response.content if b.type == "text"), "").strip()


def _ocr_page_ollama(img_bytes: bytes, config) -> str:
    import ollama
    client = ollama.Client(host=config.ollama_base_url)
    response = client.chat(
        model=config.ollama_vision_model,
        messages=[{
            "role": "user",
            "content": _OCR_PROMPT,
            "images": [img_bytes],
        }],
    )
    return response.message.content.strip()


def _parse_docx(path: Path) -> str:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)
